#!/usr/bin/env python3
"""
Post-process DOCX to apply professional table formatting:
- Header row: bold, blue background, white text
- Alternating row colors for better readability
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

def set_cell_shading(cell, hex_color):
    """Set cell background color."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)

def format_tables(doc_path, output_path=None):
    """Apply professional formatting to all tables."""
    doc = Document(doc_path)

    # Colors
    HEADER_BG = "1F4E79"  # Dark blue
    HEADER_TEXT = RGBColor(255, 255, 255)  # White
    ALT_ROW_BG = "F2F2F2"  # Light gray

    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if row_idx == 0:
                            # Header row: bold, white text
                            run.bold = True
                            run.font.color.rgb = HEADER_TEXT
                        else:
                            # Body rows: ensure black text
                            run.font.color.rgb = RGBColor(0, 0, 0)

                if row_idx == 0:
                    # Header row: blue background
                    set_cell_shading(cell, HEADER_BG)
                elif row_idx % 2 == 0:
                    # Even rows (2, 4, 6...): light gray background
                    set_cell_shading(cell, ALT_ROW_BG)

    # Save
    output = output_path or doc_path
    doc.save(output)
    print(f"Formatted tables saved to: {output}")

if __name__ == "__main__":
    input_file = sys.argv[1] if len(sys.argv) > 1 else "BRD_Intranet_Search_Analytics.docx"
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file
    format_tables(input_file, output_file)
